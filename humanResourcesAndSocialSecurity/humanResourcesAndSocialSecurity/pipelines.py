class BaikePipeline(object):
    def process_item(self, item, spider):
        return item

import json
import mimetypes
import os
import random
import time
from datetime import datetime
from markdownify import markdownify as md
from urllib.parse import urljoin, urlparse
from .xuehua.gen_md5 import get_snowflake_id

import requests
from scrapy.crawler import Crawler

import crawlab
from enum import Enum

class CrawledDataType(Enum):
    META = "meta"  # 元数据
    MASTER = "master"  # 主文档
    ATTACH = "attach"  # 附件
    HTML = "html"  # HTML源码

    # 其他需要的类型...
    @property
    def prefix(self):
        return self.value  # 返回枚举成员的值（即上面的字符串）


class TestnjPipeline(object):
    def process_item(self, item, spider):
        return item


class ProcessingResultsTracker:
    def __init__(self):
        pass

    def init_processing_results(self):
        return {
            'origin_id': None,
            'meta_record_id': None,
            'html_record_id': None,
            'master_record_id': None,
            'main_file_downloaded': False,
            'attachment_files_downloaded': [],
            'download_errors': [],
            'files_saved': {
                'meta': False,
                'html': False,
                'master': False,
                'attachments': 0
            }
        }

    def update_origin_id(self, processing_results, origin_id):
        processing_results['origin_id'] = str(origin_id) if origin_id else None

    def update_record_ids(self, processing_results, meta_id=None, html_id=None, master_id=None):
        if meta_id is not None:
            processing_results['meta_record_id'] = str(meta_id)
        if html_id is not None:
            processing_results['html_record_id'] = str(html_id)
        if master_id is not None:
            processing_results['master_record_id'] = str(master_id)

    def update_main_file_status(self, processing_results, downloaded=False, file_info=None):
        processing_results['main_file_downloaded'] = bool(downloaded)
        if file_info:
            processing_results['main_file_info'] = file_info

    def update_attachment_files(self, processing_results, downloaded_files):
        processing_results['attachment_files_downloaded'] = downloaded_files if isinstance(downloaded_files,
                                                                                           list) else []

    def add_error(self, processing_results, error_message, error_type="general"):
        import time
        error_entry = {
            'type': error_type,
            'message': str(error_message)[:200],
            'timestamp': time.time()
        }
        processing_results['download_errors'].append(error_entry)

    def mark_file_saved(self, processing_results, file_type, success=True, count=1):
        if file_type not in processing_results['files_saved']:
            return

        if file_type == 'attachments':
            if success:
                processing_results['files_saved'][file_type] += max(0, int(count))
        else:
            processing_results['files_saved'][file_type] = bool(success)

    def mark_meta_saved(self, processing_results, success=True):
        self.mark_file_saved(processing_results, 'meta', success)

    def mark_html_saved(self, processing_results, success=True):
        self.mark_file_saved(processing_results, 'html', success)

    def mark_master_saved(self, processing_results, success=True):
        self.mark_file_saved(processing_results, 'master', success)

    def mark_attachments_saved(self, processing_results, count=1):
        self.mark_file_saved(processing_results, 'attachments', True, count)

    def add_meta_error(self, processing_results, error_message):
        self.add_error(processing_results, error_message, "meta_save")
        self.mark_meta_saved(processing_results, False)

    def add_html_error(self, processing_results, error_message):
        self.add_error(processing_results, error_message, "html_save")
        self.mark_html_saved(processing_results, False)

    def add_master_error(self, processing_results, error_message):
        self.add_error(processing_results, error_message, "master_save")
        self.mark_master_saved(processing_results, False)

    def add_attachment_error(self, processing_results, error_message):
        self.add_error(processing_results, error_message, "attachment_save")

    def get_processing_statistics(self, processing_results):
        errors = processing_results.get('download_errors', [])

        stats = {
            'origin_id': processing_results.get('origin_id'),
            'total_files_attempted': 0,
            'total_files_saved': 0,
            'error_count': len(errors),
            'has_errors': len(errors) > 0,
            'success_rate': 0.0,
            'error_breakdown': self.get_error_breakdown(errors),
            'file_breakdown': {
                'main_file': processing_results.get('main_file_downloaded', False),
                'meta_file': processing_results.get('files_saved', {}).get('meta', False),
                'html_file': processing_results.get('files_saved', {}).get('html', False),
                'master_file': processing_results.get('files_saved', {}).get('master', False),
                'attachment_files': processing_results.get('files_saved', {}).get('attachments', 0)
            }
        }

        stats['total_files_attempted'] = (
                (1 if processing_results.get('main_file_downloaded') else 0) +
                len(processing_results.get('attachment_files_downloaded', [])) +
                3
        )

        stats['total_files_saved'] = (
                (1 if stats['file_breakdown']['meta_file'] else 0) +
                (1 if stats['file_breakdown']['html_file'] else 0) +
                (1 if stats['file_breakdown']['master_file'] else 0) +
                stats['file_breakdown']['attachment_files']
        )

        if stats['total_files_attempted'] > 0:
            stats['success_rate'] = (stats['total_files_saved'] / stats['total_files_attempted']) * 100

        return stats

    def get_error_breakdown(self, errors):
        breakdown = {
            'meta_save': 0,
            'html_save': 0,
            'master_save': 0,
            'attachment_save': 0,
            'general': 0,
            'critical': 0
        }

        for error in errors:
            if isinstance(error, dict):
                error_type = error.get('type', 'general')
            else:
                error_type = 'general'

            if error_type in breakdown:
                breakdown[error_type] += 1
            else:
                breakdown['general'] += 1

        return breakdown

    def get_file_status_summary(self, processing_results):
        files_saved = processing_results.get('files_saved', {})
        return {
            'meta_saved': files_saved.get('meta', False),
            'html_saved': files_saved.get('html', False),
            'master_saved': files_saved.get('master', False),
            'attachments_count': files_saved.get('attachments', 0),
            'total_core_files': sum([
                1 if files_saved.get('meta', False) else 0,
                1 if files_saved.get('html', False) else 0,
                1 if files_saved.get('master', False) else 0
            ])
        }

    def get_download_summary(self, processing_results):
        return {
            'main_file_downloaded': processing_results.get('main_file_downloaded', False),
            'attachment_files_count': len(processing_results.get('attachment_files_downloaded', [])),
            'has_main_file_info': 'main_file_info' in processing_results,
            'total_downloads': (
                    (1 if processing_results.get('main_file_downloaded', False) else 0) +
                    len(processing_results.get('attachment_files_downloaded', []))
            )
        }

    def get_errors_by_type(self, processing_results, error_type):
        errors = processing_results.get('download_errors', [])
        filtered_errors = []

        for error in errors:
            if isinstance(error, dict):
                if error.get('type') == error_type:
                    filtered_errors.append(error)
            elif error_type == 'general':
                # Backward compatibility - treat string errors as general
                filtered_errors.append({'type': 'general', 'message': str(error)})

        return filtered_errors

    def has_critical_errors(self, processing_results):
        critical_errors = self.get_errors_by_type(processing_results, 'critical')
        return len(critical_errors) > 0

    def is_processing_successful(self, processing_results):
        stats = self.get_processing_statistics(processing_results)
        return stats['success_rate'] == 100.0 and not stats['has_errors']

    def log_critical_errors(self, processing_results, spider):
        critical_errors = self.get_errors_by_type(processing_results, 'critical')
        for error in critical_errors:
            message = error.get('message', 'Unknown error') if isinstance(error, dict) else str(error)
            spider.logger.error(f"  CRITICAL: {message}")

    def log_error_details(self, processing_results, spider):
        error_types = ['meta_save', 'html_save', 'master_save', 'attachment_save', 'general']

        for error_type in error_types:
            errors = self.get_errors_by_type(processing_results, error_type)
            if errors:
                spider.logger.warning(f"  {error_type.replace('_', ' ').title()} errors ({len(errors)}):")
                for error in errors[:3]:
                    message = error.get('message', 'Unknown error') if isinstance(error, dict) else str(error)
                    spider.logger.warning(f"    - {message}")
                if len(errors) > 3:
                    spider.logger.warning(f"    ... and {len(errors) - 3} more {error_type} errors")

    def format_error_types(self, error_breakdown):
        error_parts = []
        for error_type, count in error_breakdown.items():
            if count > 0:
                error_parts.append(f"{error_type}({count})")
        return ", ".join(error_parts) if error_parts else "none"


class CustomFileStoragePipeline:
    def get_dir(self, origin_id: str, file_type: str | CrawledDataType | int | None = None) -> str:

        if file_type is None:
            return str(os.path.join(self.base_path, origin_id))

        if isinstance(file_type, CrawledDataType):
            file_type_str = file_type.prefix
        elif isinstance(file_type, int):
            try:
                crawled_type = CrawledDataType(file_type)
                file_type_str = crawled_type.prefix
            except ValueError:
                raise ValueError(f"Invalid CrawledDataType integer value: {file_type}")
        else:
            file_type_str = str(file_type)

        valid_types = {CrawledDataType.META.prefix, CrawledDataType.MASTER.prefix,
                       CrawledDataType.ATTACH.prefix, CrawledDataType.HTML.prefix}
        if file_type_str not in valid_types:
            raise ValueError(f"Unknown file_type: {file_type_str}. Valid types: {valid_types}")
        return str(os.path.join(self.base_path, origin_id, file_type_str))

    """
    自定义文件存储管道
    按照指定格式存储爬取的数据：
    - {base_path}/{origin_id}/meta/meta_{origin_id}_{record_id}.json
    - {base_path}/{origin_id}/master/master_{origin_id}_{record_id}.docx
    - {base_path}/{origin_id}/attach/attach_{origin_id}_{record_id}.xxx
    - {base_path}/{origin_id}/html/html_{origin_id}_{record_id}.html
    """

    def __init__(
            self,
            base_path=None,
            *,
            crawler: Crawler | None = None,
    ):
        if not base_path:
            raise ValueError("FILES_STORE setting is required")

        self.work_id = random.randint(1, 7)

        self.base_path = base_path
        os.makedirs(self.base_path, exist_ok=True)

    @classmethod
    def from_crawler(cls, crawler: Crawler):
        base_path = crawler.settings["FILES_STORE"]
        spider_name = crawler.spider.name
        base_path = os.path.join(base_path, spider_name)
        return cls(base_path, crawler=crawler)

    def create_tracker(self):
        return ProcessingResultsTracker()

    def _generate_numeric_id(self):
        current_time = time.time()
        timestamp_microseconds = int(current_time * 1000000)  # 微秒时间戳

        base_timestamp = str(timestamp_microseconds)[-13:]

        random_digit = random.randint(0, 9)

        numeric_id = base_timestamp + str(random_digit)

        numeric_id = numeric_id.zfill(14)

        return numeric_id

    def setup_directories(self, item):
        """创建所需的目录结构"""
        origin_id = item.get("origin_id")
        if not origin_id:
            raise ValueError("Item 'origin_id' is missing")

        # 创建所有目录
        for file_type in [CrawledDataType.META, CrawledDataType.MASTER, CrawledDataType.ATTACH, CrawledDataType.HTML]:
            os.makedirs(self.get_dir(origin_id, file_type), exist_ok=True)

    def process_item(self, item, spider):
        """处理每个item"""
        tracker = self.create_tracker()
        processing_results = tracker.init_processing_results()

        try:
            origin_id = get_snowflake_id(11, self.work_id, spider.logger)
            item["origin_id"] = origin_id
            tracker.update_origin_id(processing_results, origin_id)
            self.setup_directories(item)

            meta_record_id = origin_id
            html_record_id = get_snowflake_id(14, self.work_id, spider.logger)
            master_record_id = get_snowflake_id(15, self.work_id, spider.logger)

            # 提取内容和来源字段
            text_content, source_field = self._extract_text_content(item)
            # 将text_content存到item，供后续使用
            item['_text_content_for_doc'] = text_content

            tracker.update_record_ids(processing_results,
                                      meta_id=meta_record_id,
                                      html_id=html_record_id,
                                      master_id=master_record_id)

            main_file_info = self._download_main_file(item, master_record_id, spider)
            tracker.update_main_file_status(processing_results,
                                            downloaded=(main_file_info is not None),
                                            file_info=main_file_info)

            # 初始化URL到本地路径的映射字典
            url_to_local_path = {}

            downloaded_files = self._download_url_fields(item, master_record_id, spider)
            tracker.update_attachment_files(processing_results, downloaded_files or [])

            attachment_records = []
            if downloaded_files:
                for i, file_info in enumerate(downloaded_files):
                    if not file_info.get('is_explanation', False):
                        attach_record_id = get_snowflake_id(15, self.work_id, spider.logger)
                        attachment_records.append({
                            'file_info': file_info,
                            'record_id': attach_record_id,
                            'index': i
                        })

            # 构建URL到本地文件名的映射（只处理图片类型附件）
            image_extensions = ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')
            for i, file_info in enumerate(downloaded_files):
                # 过滤非图片附件
                if not file_info.get('extension', '').lower().endswith(image_extensions):
                    continue
                # 获取附件的原始URL
                file_url = file_info['url']
                # 获取附件的本地文件名（在save_downloaded_attachments_with_records中生成的文件名）
                # 从attachment_records中匹配对应的record_id
                if i < len(attachment_records):
                    attach_record_id = attachment_records[i]['record_id']
                    # 生成与save_downloaded_attachments_with_records一致的文件名
                    local_filename = f"attach_{origin_id}_{attach_record_id}{file_info['extension']}"
                    url_to_local_path[file_url] = local_filename

            # 将映射存入item，供后续生成MD时使用
            item['_url_to_local_image'] = url_to_local_path

            self._save_meta_data_with_tracking(item, meta_record_id, html_record_id, master_record_id,
                                               attachment_records, processing_results, tracker)
            self._save_html_data_with_tracking(item, html_record_id, processing_results, tracker)

            if downloaded_files:
                self._save_attachments_with_tracking(item, downloaded_files, attachment_records, processing_results,
                                                     tracker, spider)

            if hasattr(item, 'attachments') and item.get('attachments'):
                self._save_other_attachments_with_tracking(item, master_record_id, processing_results, tracker)

            self._save_master_document_with_tracking(item, master_record_id, main_file_info, processing_results,
                                                     tracker, source_field=source_field, url_to_local=url_to_local_path)

            self.post_process_item(item, spider, processing_results, tracker)

            # self.upload_packet_id(spider, origin_id)

        except Exception as e:
            spider.logger.error("保存文件时出错: {}".format(str(e)[:100]))
            tracker.add_error(processing_results, f"General processing error: {str(e)[:100]}", "critical")
            self.post_process_item(item, spider, processing_results, tracker)

        return item

    def _save_meta_data_with_tracking(self, item, meta_record_id, html_record_id, master_record_id, attachment_records,
                                      processing_results, tracker):
        try:
            self.save_meta_data(item, meta_record_id, html_record_id, master_record_id, attachment_records)
            tracker.mark_meta_saved(processing_results, True)
        except Exception as e:
            tracker.add_meta_error(processing_results, f"Meta save error: {str(e)[:100]}")

    def _save_html_data_with_tracking(self, item, html_record_id, processing_results, tracker):
        try:
            self.save_html_data(item, html_record_id)
            tracker.mark_html_saved(processing_results, True)
        except Exception as e:
            tracker.add_html_error(processing_results, f"HTML save error: {str(e)[:100]}")

    def _save_attachments_with_tracking(self, item, downloaded_files, attachment_records, processing_results, tracker,
                                        spider):
        try:
            self.save_downloaded_attachments_with_records(item, downloaded_files, attachment_records, spider)
            successful_attachments = len([f for f in downloaded_files if not f.get('is_explanation', False)])
            tracker.mark_attachments_saved(processing_results, successful_attachments)
        except Exception as e:
            tracker.add_attachment_error(processing_results, f"Attachments save error: {str(e)[:100]}")

    def _save_other_attachments_with_tracking(self, item, master_record_id, processing_results, tracker):
        try:
            self.save_attachments(item, master_record_id)
            attachment_count = len(item.get('attachments', []))
            tracker.mark_attachments_saved(processing_results, attachment_count)
        except Exception as e:
            tracker.add_attachment_error(processing_results, f"Other attachments save error: {str(e)[:100]}")

    def _save_master_document_with_tracking(self, item, master_record_id, main_file_info, processing_results, tracker,
                                            source_field, url_to_local=None):
        try:
            self.save_master_document(item, master_record_id, main_file_info, source_field=source_field,
                                      url_to_local=url_to_local)
            tracker.mark_master_saved(processing_results, True)
        except Exception as e:
            tracker.add_master_error(processing_results, f"Master document save error: {str(e)[:100]}")

    def post_process_item(self, item, spider, processing_results, tracker):
        stats = tracker.get_processing_statistics(processing_results)
        file_status = tracker.get_file_status_summary(processing_results)
        download_status = tracker.get_download_summary(processing_results)
        error_breakdown = stats['error_breakdown']

        if tracker.has_critical_errors(processing_results):
            spider.logger.error(
                f"CRITICAL FAILURE - Item {stats['origin_id']}: "
                f"Files: {stats['total_files_saved']}/{stats['total_files_attempted']} "
                f"({stats['success_rate']:.1f}% success) - {stats['error_count']} errors"
            )
            tracker.log_critical_errors(processing_results, spider)

        elif stats['has_errors']:
            spider.logger.warning(
                f"Item {stats['origin_id']} completed with issues: "
                f"Files: {stats['total_files_saved']}/{stats['total_files_attempted']} "
                f"({stats['success_rate']:.1f}% success) - Error types: {tracker.format_error_types(error_breakdown)}"
            )
            tracker.log_error_details(processing_results, spider)

        else:
            spider.logger.info(
                f"Item {stats['origin_id']} completed successfully: "
                f"Core files: {file_status['total_core_files']}/3, "
                f"Attachments: {file_status['attachments_count']}, "
                f"Downloads: {download_status['total_downloads']}"
            )

            crawlab.upload_packet(crawlab.UploadPacketPayload(
                packet_dir=self.get_dir(item["origin_id"]),
            ))

    def save_meta_data(self, item, meta_record_id, html_record_id, master_record_id, attachment_records=None):
        """保存meta数据为JSON文件"""
        origin_id = item.get("origin_id")
        filename = "meta_{}_{}.json".format(origin_id, origin_id)
        filepath = os.path.join(self.get_dir(item["origin_id"], CrawledDataType.META), filename)

        # 从URL提取域名作为网站名称
        web_site = self._extract_domain_name(item.get('_response_url', ''))

        # 获取当前采集时间
        crawl_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S.%f")

        # 获取发布时间
        publish_time = self._extract_publish_time(item)

        # 获取标签信息
        tags = self._extract_tags(item)

        # 构建附件列表
        attachment_name_list = item.get('attachment_name', [])
        attachment_list, attachment = self._build_attachment_list_with_records(attachment_records or [],
                                                                               attachment_name_list or [])

        # 直接使用爬取的文本内容
        content, _ = self._extract_text_content(item)

        excluded_fields = [
            'title', 'publish_time', 'content', 'md_content', 'tags', 'origin_id', 'attachment', 'attachment_name'
        ]
        remaining_fields = {}
        for field_key, field_value in item.items():
            if (field_key not in excluded_fields
                    and not field_key.startswith('_')):
                # 完全保留原始格式（列表、字符串、数字等不做任何处理）
                remaining_fields[field_key] = field_value

        # 构建新格式的meta数据
        meta_data_item = {
            "web_site": web_site,
            "src_url": item.get("_response_url", ""),
            "publish_time": publish_time if publish_time else self.extract_date(crawl_time),
            "crawl_time": crawl_time,
            "file_id": int(master_record_id),  # 使用主文档的record_id
            "html_id": int(html_record_id),  # 使用html文件的record_id
            "title": self._extract_title(item),
            "attachment_list": attachment_list,
            "content": content,
            "tags": tags,
            "extend": {
                'attachment': attachment,
                **remaining_fields
            },  # TODO! 固定为空对象，后续支持自定义
        }

        meta_data_item_info = {
            "src_url": item.get("_response_url", ""),
            "file_id": int(master_record_id),  # 使用主文档的record_id
            "publish_time": publish_time if publish_time else self.extract_date(crawl_time),
            "title": self._extract_title(item),
            "content": content ,
            "extend": {
                'attachment': attachment,
            },

        }

        crawlab.save_item(meta_data_item_info)

        final_data = {
            "runner": "",  # 运行器标识，暂时为空
            "meta_data": [meta_data_item]
        }

        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(final_data, f, ensure_ascii=False, indent=2)

    def _download_main_file(self, item, record_id, spider):
        """下载主文件（file_category='main'的URL字段）"""

        # 获取字段分类信息
        field_categories = item.get('_field_categories', {})

        # 查找第一个file_category='main'的URL字段
        for field_name, field_value in item.items():
            # 跳过内部字段
            if field_name.startswith('_'):
                continue

            # 只处理file_category='main'的字段
            if field_categories.get(field_name, '') != 'main':
                continue

            # 获取URL
            urls = self._get_urls_from_value(field_value)
            if not urls:
                continue

            url = urls[0]  # 取第一个URL

            # 处理相对路径
            if not url.startswith(('http://', 'https://')):
                base_url = item.get('_response_url', '')
                if base_url:
                    from urllib.parse import urljoin
                    url = urljoin(base_url, url)
                else:
                    continue

            try:
                # 简单下载
                import requests
                response = requests.get(url, timeout=30)
                response.raise_for_status()

                # 获取文件扩展名
                file_extension = self._get_file_extension_from_url(url, response.headers)

                # 下载主文件成功

                return {
                    'data': response.content,
                    'url': url,
                    'extension': file_extension
                }

            except Exception as e:
                spider.logger.error("主文件下载失败: {}".format(str(e)[:100]))
                continue

        return None

    def _extract_domain_name(self, url):
        """从URL提取域名"""
        if not url:
            return ""

        try:
            from urllib.parse import urlparse
            parsed = urlparse(url)
            domain = parsed.netloc.lower()

            # 直接返回域名，保持原始格式
            return domain if domain else ""

        except Exception:
            return ""

    def _extract_title(self, item):
        # 只从item的字段中提取，没有就为空
        title_fields = ['title', 'Title', '标题']

        for field in title_fields:
            if field in item and item[field]:
                title = item[field]
                if isinstance(title, list):
                    title = ' '.join(str(t) for t in title if t)
                return str(title).strip()

        return ""

    def _extract_tags(self, item):
        # 只从item的字段中提取，没有就为空
        tags_fields = ['tags', '标签']
        tags = []

        for field in tags_fields:
            if field in item and item[field]:
                tags_value = item[field]
                if isinstance(tags_value, list):
                    tags = list(dict.fromkeys([tag.strip() for tag in tags_value if tag.strip()]))
                elif isinstance(tags_value, str):
                    tags.append(tags_value)

        return tags

    def _extract_publish_time(self, item):
        # 只从item的字段中提取，没有就为空
        time_fields = ['publish_time', 'date', 'time', '发布时间', '时间']

        for field in time_fields:
            if field in item and item[field]:
                time_value = item[field]
                if isinstance(time_value, list):
                    time_value = time_value[0] if time_value else ""

                time_str = self.extract_date(str(time_value))
                if time_str:
                    return time_str

        return ""

    def extract_date(self, text):
        """从文本中提取日期（支持多种格式）"""
        import re
        # 匹配：YYYY-MM-DD、YYYY/MM/DD、YYYY年MM月DD日、YYYY.MM.DD、YYYYMMDD
        match = re.search(
            r'(\d{4})[-/年.](\d{1,2})[-/月.](\d{1,2})[日]?|(\d{4})(\d{2})(\d{2})',
            text
        )
        if match:
            # 处理第一种格式（YYYY-MM-DD等）
            if match.group(1):
                year, month, day = match.group(1), match.group(2), match.group(3)
            # 处理YYYYMMDD格式
            else:
                year, month, day = match.group(4), match.group(5), match.group(6)
            return f"{year}-{month.zfill(2)}-{day.zfill(2)}"  # 补零并统一格式
        return ''  # 如果没有匹配到任何格式，返回空字符串

    def _extract_text_content(self, item):
        # 优先使用预处理的文本内容
        content = ''
        source_field = None
        if not content:
            # 尝试从其他字段获取内容
            content_fields = ['content', 'text', '内容', 'body', 'article']
            for field in content_fields:
                if field in item and item[field]:
                    content = item[field]
                    source_field = field
                    if isinstance(content, list):
                        content = '\n'.join(str(c) for c in content if c)
                    break

        if not content:
            # 尝试从其他字段获取内容
            content_fields = ['md_content']
            for field in content_fields:
                if field in item and item[field]:
                    content = item[field]
                    source_field = field
                    if isinstance(content, list):
                        content = '\n'.join(str(c) for c in content if c)
                    content = md(content)
                    break

        if isinstance(content, str):
            processed_content = content.strip()
        else:
            processed_content = str(content) if content else ""

            # 返回内容和来源字段（元组形式）
        return processed_content, source_field

    def _build_attachment_list_with_records(self, attachment_records, attachment_names):

        attachment_list = []
        attachment = []

        for idx, record in enumerate(attachment_records):
            file_info = record['file_info']
            attach_record_id = record['record_id']

            attach_name = ""
            url = file_info['url']
            attachment.append(url)
            if idx < len(attachment_names) and attachment_names[idx].strip():
                attach_name = attachment_names[idx].strip()

            else:
                attach_name = url.split('/')[-1] if '/' in url else url
                if '?' in attach_name:
                    attach_name = attach_name.split('?')[0]

            if not attach_name:
                attach_name = "attachment_{}_{}{}".format(
                    record['index'] + 1,
                    attach_record_id,
                    file_info.get('extension', '')
                )

            attachment_list.append({
                "attach_id": int(attach_record_id),  # 使用独立生成的record_id
                "attach_name": attach_name
            })

        return attachment_list, attachment

    def save_downloaded_attachments_with_records(
            self, item, downloaded_files, attachment_records, spider
    ):
        # 创建记录映射
        record_map = {}
        for record in attachment_records:
            record_map[record['index']] = record['record_id']

        origin_id = item.get("origin_id")
        for i, file_info in enumerate(downloaded_files):
            try:
                if file_info.get('is_explanation', False):
                    attach_record_id = get_snowflake_id(15, self.work_id, spider.logger)
                    filename = "download_failed_explanation_{}_{}{}".format(
                        origin_id, attach_record_id, file_info["extension"]
                    )
                else:
                    if i in record_map:
                        attach_record_id = record_map[i]
                        filename = "attach_{}_{}{}".format(
                            origin_id, attach_record_id, file_info["extension"]
                        )
                    else:
                        continue

                filepath = os.path.join(
                    self.get_dir(item["origin_id"], CrawledDataType.ATTACH), filename
                )

                # 确保目录存在
                os.makedirs(self.get_dir(item["origin_id"], CrawledDataType.ATTACH), exist_ok=True)

                # 保存文件
                with open(filepath, 'wb') as f:
                    f.write(file_info['data'])

            except Exception:
                pass  # 静默处理错误，避免中断主流程

    def save_html_data(self, item, record_id):
        origin_id = item.get("origin_id")
        filename = "html_{}_{}.html".format(origin_id, record_id)
        filepath = os.path.join(self.get_dir(item["origin_id"], CrawledDataType.HTML), filename)

        # 获取页面源码
        if '_response_html' in item:
            # 如果item中保存了页面源码，使用原始HTML
            html_content = item['_response_html']
        elif '_response_url' in item:
            # 如果只有URL，创建一个包含URL信息的HTML
            html_content = """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>页面源码 - {}</title>
</head>
<body>
    <h1>页面信息</h1>
    <p><strong>URL:</strong> {}</p>
    <p><strong>说明:</strong> 原始HTML源码未能获取，显示页面基本信息</p>

    <h2>提取的数据:</h2>
    <div class="item-data">
        {}
    </div>

    <div class="meta-info">
        <p>Origin ID: {}</p>
        <p>Record ID: {}</p>
        <p>Created: {}</p>
    </div>
</body>
</html>""".format(
                record_id,
                item.get("_response_url", "未知URL"),
                self._format_item_data(item),
                item.get("origin_id"),
                record_id,
                datetime.now().isoformat(),
            )
        else:
            # 降级方案：从item数据生成HTML（保持向后兼容）
            html_content = self._create_html_content(item, record_id)

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)

    def save_attachments(self, item, record_id):
        attachments = item.get('attachments', [])
        for i, attachment in enumerate(attachments):
            # 根据附件类型确定文件扩展名
            if isinstance(attachment, dict) and 'url' in attachment:
                url = attachment['url']
                file_ext = self._get_file_extension_from_url(url, {})
            else:
                file_ext = '.txt'  # 默认扩展名
            origin_id = item.get("origin_id")
            filename = "attach_{}_{}{}{}".format(origin_id, record_id, i + 1, file_ext)
            filepath = os.path.join(self.get_dir(item["origin_id"], CrawledDataType.ATTACH), filename)

            # 保存附件内容
            with open(filepath, 'wb') as f:
                if isinstance(attachment, bytes):
                    f.write(attachment)
                else:
                    f.write(str(attachment).encode('utf-8'))

    def save_master_document(self, item, record_id, main_file_info=None, source_field=None, url_to_local=None):
        # 如果有下载的主文件，直接保存
        if main_file_info and main_file_info.get('data'):
            origin_id = item.get("origin_id")
            file_extension = main_file_info.get('extension', '.bin')
            filename = "master_{}_{}{}".format(origin_id, record_id, file_extension)
            filepath = os.path.join(self.get_dir(item["origin_id"], CrawledDataType.MASTER), filename)

            with open(filepath, 'wb') as f:
                f.write(main_file_info['data'])
            return

        # 没有下载的主文件，基于text内容生成docx
        title = item.get('title', '').strip()
        text_content = item.get('_text_content_for_doc', '')
        if not text_content:
            return

        origin_id = item.get("origin_id")

        if source_field == 'md_content':
            # Markdown文件逻辑：直接写入文本内容
            filename = "master_{}_{}.md".format(origin_id, record_id)
            filepath = os.path.join(self.get_dir(item["origin_id"], CrawledDataType.MASTER), filename)

            if url_to_local and isinstance(text_content, str):
                import re
                # 匹配Markdown图片格式：![alt](url) 或 ![alt](url "title")
                pattern = r'!\[(.*?)\]\(\s*([^\s\)]+)\s*(?:["\'](.*?)["\'])?\s*\)'

                def replace_match(match):
                    alt_text = match.group(1)
                    img_url = match.group(2).strip()  # 原始图片URL
                    title = match.group(3) or ''

                    # 获取页面的基础URL（爬取该页面时的原始URL，用于拼接相对路径）
                    base_url = item.get('_response_url', '')
                    if base_url and img_url:
                        # 判断img_url是否为相对路径（无http/https协议头）
                        if not img_url.startswith(('http://', 'https://')):
                            # 用base_url拼接成完整绝对路径
                            img_url = urljoin(base_url, img_url)

                    # 查找URL对应的本地文件名
                    local_filename = url_to_local.get(img_url)
                    if not local_filename:
                        # 未找到对应附件，保留原始URL
                        return match.group(0)

                    # 计算相对路径：master目录 -> 上级目录 -> attach目录 -> 文件名
                    relative_path = f"../attach/{local_filename}"

                    # 重构图片标签（保留alt和title）
                    if title:
                        return f'![{alt_text}]({relative_path} "{title}")'
                    else:
                        return f'![{alt_text}]({relative_path})'

                # 执行替换
                text_content = re.sub(pattern, replace_match, text_content)

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(md(f'<h1>{title}</h1>') + '\n\n' + text_content)
            return

        filename = "master_{}_{}.docx".format(origin_id, record_id)
        filepath = os.path.join(self.get_dir(item["origin_id"], CrawledDataType.MASTER), filename)

        try:
            self._create_simple_docx(filepath, title, text_content, item, record_id)
        except ImportError:
            self._create_simple_text(filepath, text_content, item, record_id)
        except Exception:
            self._create_simple_text(filepath, text_content, item, record_id)

    def _create_simple_docx(self, filepath, title, text_content, item, record_id, head_font_name="方正小标宋简体",
                            par_font_name="仿宋", text_indent=0.5, text_spacing=1.5):
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT

        doc = Document()
        # 设置标题样式
        head = doc.add_heading(title, 1)
        """标题文本写入"""
        head_text = head.runs[0]
        head_text.font.name = head_font_name
        head_text.font.size = Pt(18)
        head_text.font.bold = True
        """设置居中对齐 CENTER LEFT RIGHT"""
        head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = head_text._element.rPr.rFonts
        r.set(qn('w:eastAsia'), head_font_name)

        for line in text_content.split('\n'):
            # 去除每行末尾的空白字符（包括换行符和空格）
            stripped_line = line.strip()
            if stripped_line:  # 检查行是否为空（只包含空白字符）
                # 创建一个新段落并设置缩进
                par = doc.add_paragraph(stripped_line)
                par.paragraph_format.first_line_indent = Inches(text_indent)
                paragraph_text = par.runs[0]
                paragraph_text.font.name = par_font_name
                paragraph_text.font.size = Pt(16)
                r = paragraph_text._element.rPr.rFonts
                r.set(qn('w:eastAsia'), par_font_name)
        for paragraph in doc.paragraphs:
            """设置行间距"""
            paragraph.paragraph_format.line_spacing = text_spacing
        # 保存Word文档
        doc.save(filepath)

    def _create_simple_text(self, filepath, text_content, item, record_id):
        # 将.docx扩展名改为.txt
        txt_filepath = filepath.replace('.docx', '.txt')

        with open(txt_filepath, 'w', encoding='utf-8') as f:
            f.write("主文档 - {}\n".format(record_id))
            f.write("=" * 50 + "\n\n")

            # 写入页面信息
            if '_response_url' in item:
                f.write("页面信息:\n")
                f.write("-" * 20 + "\n")
                f.write("数据来源: {}\n\n".format(item.get('_response_url', '未知')))

            # 写入文本内容
            f.write("提取的文本内容:\n")
            f.write("-" * 20 + "\n")
            f.write(text_content)

            # 写入文档信息
            f.write("\n\n文档信息:\n")
            f.write("-" * 20 + "\n")
            origin_id = item.get("origin_id")
            f.write("Origin ID: {}\n".format(origin_id))
            f.write("Record ID: {}\n".format(record_id))
            f.write("创建时间: {}\n".format(datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
            f.write("数据来源: {}\n".format(item.get('_response_url', '未知')))

    def _format_item_data(self, item):
        html_parts = []
        for key, value in item.items():
            # 跳过内部字段
            if key.startswith('_'):
                continue
            if isinstance(value, (list, tuple)):
                value_str = ', '.join(str(v) for v in value)
            else:
                value_str = str(value)
            html_parts.append("<p><strong>{}:</strong> {}</p>".format(key, value_str))
        return '\n        '.join(html_parts)

    def _create_html_content(self, item, record_id):
        """创建HTML内容"""
        html_parts = self._format_item_data(item)

        return """<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Scraped Data - {}</title>
</head>
<body>
    <h1>爬取数据 - {}</h1>
    <div class="item-data">
        {}
    </div>
    <div class="meta-info">
        <p>Origin ID: {}</p>
        <p>Record ID: {}</p>
        <p>Created: {}</p>
    </div>
</body>
</html>""".format(
            record_id,
            record_id,
            html_parts,
            item.get("origin_id"),
            record_id,
            datetime.now().isoformat(),
        )

    def _download_url_fields(self, item, record_id, spider):
        """检测并下载URL类型字段的文件"""
        downloaded_files = []

        # 开始检测URL字段

        # 获取字段分类信息
        field_categories = item.get('_field_categories', {})
        # 获取字段分类信息

        # 遍历item的所有字段，查找包含http链接的字段
        for field_name, field_value in item.items():
            # 跳过内部字段
            if field_name.startswith('_'):
                continue

            # 检查字段的file_category，只下载标记为'attachment'的字段
            field_category = field_categories.get(field_name, '')
            if field_category != 'attachment':
                # 跳过非attachment类型的字段
                continue

            # 检查attachment字段

            # 获取URL列表
            urls = self._get_urls_from_value(field_value)
            urls = list(dict.fromkeys([u for u in urls if isinstance(u, str)]))

            if urls:
                spider.logger.debug("发现{}个待下载文件".format(len(urls)))

            for url in urls:
                try:
                    # 处理相对路径，转换为绝对URL
                    if not url.startswith(('http://', 'https://')):
                        # 从item中获取当前页面URL作为base
                        base_url = item.get('_response_url', '')
                        if base_url:
                            # 使用urljoin处理相对路径
                            url = urljoin(base_url, url)
                            # 相对路径转换为绝对URL
                        else:
                            # 无法获取base URL，跳过相对路径
                            continue

                    spider.logger.debug("正在下载: {}".format(url))

                    # 下载文件，设置完整的浏览器headers
                    from fake_useragent import UserAgent
                    ua = UserAgent()
                    u_a = ua.random
                    headers = {
                        'User-Agent': u_a,
                        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
                        'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
                        'Accept-Encoding': 'gzip, deflate, br',
                        'DNT': '1',
                        'Connection': 'keep-alive',
                        'Upgrade-Insecure-Requests': '1',
                        'Sec-Fetch-Dest': 'document',
                        'Sec-Fetch-Mode': 'navigate',
                        'Sec-Fetch-Site': 'same-origin',
                        'Cache-Control': 'max-age=0'
                    }

                    # 添加Referer头，模拟从原页面点击下载
                    if '_response_url' in item:
                        headers['Referer'] = item['_response_url']

                    # 使用session保持连接
                    session = requests.Session()
                    session.headers.update(headers)

                    # 先访问原页面建立正常的session（模拟用户正常浏览行为）
                    if '_response_url' in item:
                        try:
                            session.get(item['_response_url'], timeout=10)
                            # 已访问原页面建立session
                        except:
                            pass  # 忽略原页面访问失败

                    response = session.get(url, timeout=10, allow_redirects=True)
                    response.raise_for_status()

                    # 检查返回内容是否为正常文件
                    content_type = response.headers.get('content-type', '').lower()
                    content_length = response.headers.get('content-length', '0')

                    # 检查是否为HTML/JSP页面（通常是错误页面或验证页面）
                    is_anti_crawl = False
                    anti_crawl_reason = ""

                    if 'text/html' in content_type:
                        try:
                            response_text = response.content.decode('utf-8', errors='ignore')
                            # 检查内容是否包含HTML标签或JSP错误页面特征
                            html_indicators = ['<html', '<body', '<head', '<!DOCTYPE', 'jsp', '验证码', 'captcha',
                                               'codeValue', '附件下载', '请输入', 'error', '错误', '<script>', '<div>']
                            if any(indicator in response_text.lower() for indicator in html_indicators):
                                is_anti_crawl = True
                                anti_crawl_reason = "服务器返回HTML/JSP页面而非文件，可能是验证码页面或反爬虫机制"
                                # 检测到反爬虫机制
                        except:
                            pass

                    # 检查文件大小是否异常小
                    response_size = len(response.content)
                    if not is_anti_crawl and response_size < 1024:  # 小于1KB
                        is_anti_crawl = True
                        anti_crawl_reason = "文件大小异常小（{}字节），可能是错误页面或反爬虫响应".format(response_size)
                        # 文件大小异常，可能是错误页面

                    # 特别检查JSP下载链接
                    if not is_anti_crawl and (url.endswith('.jsp') or 'download.jsp' in url):
                        try:
                            response_text = response.content.decode('utf-8', errors='ignore')
                            if '<' in response_text and '>' in response_text:  # 包含HTML标签
                                is_anti_crawl = True
                                anti_crawl_reason = "JSP下载链接返回HTML内容，需要浏览器JavaScript执行或验证码验证"
                                # JSP页面包含HTML内容
                        except:
                            pass

                    # 如果检测到反爬虫，生成说明文件
                    if is_anti_crawl:
                        explanation_content = (
                            "文件下载失败说明\n"
                            "===================\n\n"
                            "原始链接: {}\n"
                            "检测时间: {}\n"
                            "文件大小: {} 字节\n"
                            "Content-Type: {}\n\n"
                            "失败原因:\n"
                            "{}\n\n"
                            "解决建议:\n"
                            "1. 使用浏览器手动下载该文件\n"
                            "2. 该网站可能设置了反爬虫机制,需要验证码或JavaScript验证\n"
                            "3. 可以尝试使用浏览器自动化工具进行下载\n"
                            "4. 检查是否需要特定的登录状态或权限\n\n"
                            "技术说明:\n"
                            "此文件无法通过程序自动下载,服务器可能检测到了自动化请求并返回了验证页面.\n"
                            "建议手动访问原始链接进行下载.\n"
                        ).format(url, datetime.now().strftime('%Y-%m-%d %H:%M:%S'), response_size, content_type,
                                 anti_crawl_reason)

                        downloaded_files.append({
                            'data': explanation_content.encode('utf-8'),
                            'extension': '.txt',
                            'url': url,
                            'is_explanation': True
                        })
                        # 生成反爬虫说明文件
                        continue

                    # 获取文件扩展名
                    file_ext = self._get_file_extension_from_url(url, response.headers)

                    downloaded_files.append({
                        'data': response.content,
                        'extension': file_ext,
                        'url': url
                    })
                    spider.logger.debug("下载完成: {} ({}KB)".format(url.split('/')[-1], len(response.content) // 1024))

                except Exception as e:
                    spider.logger.error("文件下载失败: {}".format(str(e)[:100]))

        if downloaded_files:
            spider.logger.info("文件下载任务完成，共下载{}个文件".format(len(downloaded_files)))
        return downloaded_files

    def _get_urls_from_value(self, value):
        """从字段值中提取所有HTTP URL"""
        urls = []

        if isinstance(value, str):
            urls.extend(value.split(' '))
        elif isinstance(value, (list, tuple)):
            for item in value:
                if isinstance(item, str):
                    urls.append(item)

        return urls

    def _get_file_extension_from_url(self, url, headers):
        """从URL和响应头获取文件扩展名"""
        # 常见文件扩展名白名单
        valid_extensions = ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
                            '.txt', '.zip', '.rar', '.7z', '.tar', '.gz', '.jpg', '.jpeg',
                            '.png', '.gif', '.bmp', '.mp3', '.mp4', '.avi', '.mov', '.wmv',
                            '.webp', '.ofd']
        # 首先尝试从URL路径获取扩展名
        parsed_url = urlparse(url)
        path = parsed_url.path
        if path and '.' in path:
            ext = os.path.splitext(path)[1].lower()
            if ext in valid_extensions:
                return ext

        # 尝试从Content-Type获取扩展名
        content_type = headers.get('content-type', '').lower().split(';')[0]
        content_type_mapping = {
            'application/pdf': '.pdf',
            'application/msword': '.doc',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'application/vnd.ms-excel': '.xls',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
            'application/vnd.ms-powerpoint': '.ppt',
            'application/vnd.openxmlformats-officedocument.presentationml.presentation': '.pptx',
            'application/zip': '.zip',
            'application/x-rar-compressed': '.rar',
            'application/x-7z-compressed': '.7z',
            'text/plain': '.txt',
            'image/jpeg': '.jpg',
            'image/png': '.png',
            'image/gif': '.gif',
            'image/bmp': '.bmp',
            'audio/mpeg': '.mp3',
            'video/mp4': '.mp4',
            'video/avi': '.avi',
            'video/quicktime': '.mov'
        }

        if content_type in content_type_mapping:
            return content_type_mapping[content_type]

        # 尝试从Content-Disposition获取扩展名
        content_disposition = headers.get('content-disposition', '').lower()
        for ext in valid_extensions:
            if ext in content_disposition:
                return ext

        # 尝试从filename获取扩展名
        filename = headers.get('filename', '').lower()
        for ext in valid_extensions:
            if ext in filename:
                return ext

        # 使用mimetypes模块作为备选
        if content_type:
            ext = mimetypes.guess_extension(content_type)
            if ext:
                return ext

        # 默认扩展名
        return '.bin'

    def close_spider(self, spider):
        """爬虫关闭时上传文件"""

